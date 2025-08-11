"""
Document Loader for Custom Data Integration
Supports PDF, DOCX, TXT, CSV, Excel and directory-based loading
"""

import os
import pandas as pd
import numpy as np
from typing import List, Dict, Optional, Tuple
from pathlib import Path
import warnings
# Document objects for RAGAS compatibility
try:
    from langchain_core.documents import Document
    LANGCHAIN_AVAILABLE = True
except ImportError:
    try:
        from langchain.schema import Document
        LANGCHAIN_AVAILABLE = True
    except ImportError:
        LANGCHAIN_AVAILABLE = False
        print("⚠️ Warning: langchain not available, Document objects will not be created")
        
        # Fallback Document class
        class Document:
            def __init__(self, page_content: str, metadata: dict = None):
                self.page_content = page_content
                self.metadata = metadata or {}

warnings.filterwarnings('ignore')

# PDF processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    print("⚠️  PDF support not available. Install with: pip install PyPDF2 pdfplumber")
    PDF_AVAILABLE = False

# Word document processing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    print("⚠️  DOCX support not available. Install with: pip install python-docx")
    DOCX_AVAILABLE = False

# Advanced text processing
try:
    from langdetect import detect, detect_langs
    import langdetect
    LANGDETECT_AVAILABLE = True
except ImportError:
    print("⚠️  Language detection not available. Install with: pip install langdetect")
    LANGDETECT_AVAILABLE = False

# Topic extraction
try:
    from keybert import KeyBERT
    KEYBERT_AVAILABLE = True
except ImportError:
    KEYBERT_AVAILABLE = False

try:
    import yake
    YAKE_AVAILABLE = True
except ImportError:
    YAKE_AVAILABLE = False

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.decomposition import LatentDirichletAllocation
    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False

class DocumentLoader:
    """Comprehensive document loader supporting multiple formats"""
    
    def __init__(self, config: Dict):
        """Initialize document loader with configuration"""
        self.config = config
        self.custom_config = config.get('custom_data', {})
        self.processing_config = self.custom_config.get('processing', {})
        self.topic_config = self.custom_config.get('topic_extraction', {})
        
        # Initialize topic extraction models with proper offline configuration
        self.keybert_model = self._initialize_keybert_offline()
        self.yake_extractor = yake.KeywordExtractor(
            lan="en", n=3, dedupLim=0.7, top=10
        ) if YAKE_AVAILABLE else None
        
        # Document storage
        self.documents = []
        self.document_metadata = []
    
    def _get_best_device(self):
        """Get the best available device for KeyBERT based on configuration"""
        try:
            # First check testset_generation > knowledge_graph > sentence_model_config
            kg_config = self.config.get('testset_generation', {}).get('knowledge_graph', {})
            sentence_config = kg_config.get('sentence_model_config', {})
            config_device = sentence_config.get('device', None)
            
            # Fallback to LLM config if not found in testset generation
            if not config_device:
                keybert_config = self.config.get('llm', {}).get('keyword_extraction', {}).get('keybert', {})
                config_device = keybert_config.get('device', 'cpu')
            
            # If config specifies CPU, use it
            if config_device == 'cpu':
                print("💻 Using CPU device as configured")
                return 'cpu'
            
            # Parse device configuration (could be "cuda:0, cuda:1" or "cuda:0")
            if isinstance(config_device, str) and 'cuda' in config_device:
                # Extract first CUDA device from comma-separated list
                cuda_devices = [d.strip() for d in config_device.split(',')]
                for device in cuda_devices:
                    if device.startswith('cuda'):
                        try:
                            import torch
                            if torch.cuda.is_available():
                                device_id = device.split(':')[1] if ':' in device else '0'
                                if int(device_id) < torch.cuda.device_count():
                                    print(f"🎮 Using GPU device: {device}")
                                    return device
                        except Exception as e:
                            print(f"⚠️  GPU device {device} not available: {e}")
                            continue
            
            # Fallback to CPU if no valid GPU found
            print("💻 Falling back to CPU device")
            return 'cpu'
            
        except Exception as e:
            print(f"⚠️  Device detection error: {e}")
            return 'cpu'

    def _initialize_keybert_offline(self):
        """Initialize KeyBERT with offline mode and GPU support based on configuration"""
        if not KEYBERT_AVAILABLE:
            print("⚠️  KeyBERT not available - skipping keyword extraction")
            return None
            
        try:
            # Get KeyBERT configuration from config
            keybert_config = self.config.get('llm', {}).get('keyword_extraction', {}).get('keybert', {})
            
            # Check if KeyBERT is enabled in config
            if not keybert_config.get('enabled', True):
                print("⚠️  KeyBERT disabled in configuration")
                return None
            
            # Get model name from config
            model_name = keybert_config.get('model', 'all-MiniLM-L6-v2')
            print(f"🔧 Initializing KeyBERT with model: {model_name}")
            
            # Get cache directory from config (matching offline model manager)
            cache_dir = self.config.get('advanced', {}).get('caching', {}).get('cache_dir', './cache')
            sentence_cache_dir = os.path.join(cache_dir, 'sentence_transformers')
            
            # Get best available device for GPU acceleration
            device = self._get_best_device()
            
            # Try to initialize with offline mode first
            try:
                from sentence_transformers import SentenceTransformer
                print(f"🔄 Loading sentence transformer: {model_name} on {device}")
                
                # Initialize SentenceTransformer with offline mode and GPU support
                sentence_model = SentenceTransformer(
                    model_name,
                    cache_folder=sentence_cache_dir,
                    local_files_only=True,
                    device=device
                )
                
                # Initialize KeyBERT with the offline model
                keybert_model = KeyBERT(model=sentence_model)
                print(f"✅ KeyBERT initialized with offline model: {model_name} on {device}")
                return keybert_model
                
            except Exception as e:
                print(f"⚠️  KeyBERT offline initialization failed on {device}: {e}")
                
                # Try CPU fallback if GPU failed
                if device != 'cpu':
                    print("🔄 Attempting CPU fallback with local_files_only...")
                    try:
                        from sentence_transformers import SentenceTransformer
                        sentence_model = SentenceTransformer(
                            model_name,
                            cache_folder=sentence_cache_dir,
                            local_files_only=True,
                            device='cpu'
                        )
                        keybert_model = KeyBERT(model=sentence_model)
                        print(f"✅ KeyBERT initialized with offline model on CPU: {model_name}")
                        return keybert_model
                    except Exception as e_cpu:
                        print(f"⚠️  CPU fallback with offline mode also failed: {e_cpu}")
                
                print("🔄 Attempting fallback without local_files_only...")
                
                # Fallback: Try without local_files_only but with cache
                try:
                    from sentence_transformers import SentenceTransformer
                    fallback_device = 'cpu' if device != 'cpu' else device
                    sentence_model = SentenceTransformer(
                        model_name,
                        cache_folder=sentence_cache_dir,
                        device=fallback_device
                    )
                    keybert_model = KeyBERT(model=sentence_model)
                    print(f"✅ KeyBERT initialized with fallback method: {model_name} on {fallback_device}")
                    return keybert_model
                    
                except Exception as e2:
                    print(f"⚠️  KeyBERT fallback initialization also failed: {e2}")
                    print("🔄 Final fallback: Disabling KeyBERT - continuing without keyword extraction")
                    return None
                    
        except Exception as e:
            print(f"⚠️  KeyBERT configuration error: {e}")
            print("🔄 Disabling KeyBERT - continuing without keyword extraction")
            return None
        
    def load_all_documents(self) -> Tuple[List[Document], List[Dict]]:
        """Load all documents from configured sources"""
        print("📂 Loading documents from configured sources...")
        
        # Check if we're using pipeline config structure
        if 'data_sources' in self.config:
            pipeline_data_sources = self.config['data_sources']
            input_type = pipeline_data_sources.get('input_type', 'documents')
            
            if input_type == 'csv':
                print("📊 Using CSV input mode from pipeline config")
                # Load CSV files from pipeline config
                csv_config = pipeline_data_sources.get('csv', {})
                csv_files = csv_config.get('csv_files', [])
                
                for csv_file in csv_files:
                    if os.path.exists(csv_file):
                        print(f"  📊 Loading CSV: {csv_file}")
                        docs, metadata = self.load_csv_content(csv_file)
                        self.documents.extend(docs)
                        self.document_metadata.extend(metadata)
                    else:
                        print(f"  ⚠️  CSV file not found: {csv_file}")
                
                # Process and return
                if self.documents:
                    print(f"📋 Processing {len(self.documents)} raw documents...")
                    processed_docs, processed_metadata = self.process_documents(
                        self.documents, self.document_metadata
                    )
                    print(f"✅ Final document count: {len(processed_docs)}")
                    
                    # Convert to Document objects for RAGAS compatibility
                    if LANGCHAIN_AVAILABLE:
                        document_objects = []
                        for doc_content, metadata in zip(processed_docs, processed_metadata):
                            doc_obj = Document(
                                page_content=doc_content,
                                metadata=metadata
                            )
                            document_objects.append(doc_obj)
                        print(f"📄 Created {len(document_objects)} Document objects for RAGAS")
                        return document_objects, processed_metadata
                    else:
                        print("⚠️ langchain not available, returning strings")
                        return processed_docs, processed_metadata
                else:
                    print("⚠️  No CSV documents loaded. Using default document corpus.")
                    default_docs, default_metadata = self.get_default_documents()
                    
                    # Convert default documents to Document objects
                    if LANGCHAIN_AVAILABLE and isinstance(default_docs[0], str):
                        document_objects = []
                        for doc_content, metadata in zip(default_docs, default_metadata):
                            doc_obj = Document(
                                page_content=doc_content,
                                metadata=metadata
                            )
                            document_objects.append(doc_obj)
                        return document_objects, default_metadata
                    else:
                        return default_docs, default_metadata
        
        # Original document loading logic
        data_sources = self.custom_config.get('data_sources', {})
        
        # Load PDF files
        pdf_files = data_sources.get('pdf_files', [])
        for pdf_file in pdf_files:
            if os.path.exists(pdf_file):
                print(f"  📄 Loading PDF: {pdf_file}")
                docs, metadata = self.load_pdf(pdf_file)
                self.documents.extend(docs)
                self.document_metadata.extend(metadata)
            else:
                print(f"  ⚠️  PDF not found: {pdf_file}")
        
        # Load text files
        text_files = data_sources.get('text_files', [])
        for text_file in text_files:
            if os.path.exists(text_file):
                print(f"  📝 Loading text: {text_file}")
                docs, metadata = self.load_text_file(text_file)
                self.documents.extend(docs)
                self.document_metadata.extend(metadata)
            else:
                print(f"  ⚠️  Text file not found: {text_file}")
        
        # Load Word documents
        word_files = data_sources.get('word_files', [])
        for word_file in word_files:
            if os.path.exists(word_file):
                print(f"  📄 Loading DOCX: {word_file}")
                docs, metadata = self.load_docx(word_file)
                self.documents.extend(docs)
                self.document_metadata.extend(metadata)
            else:
                print(f"  ⚠️  DOCX file not found: {word_file}")
        
        # Load from directories
        directories = data_sources.get('directories', [])
        for directory in directories:
            if os.path.exists(directory):
                print(f"  📁 Loading from directory: {directory}")
                docs, metadata = self.load_directory(directory)
                self.documents.extend(docs)
                self.document_metadata.extend(metadata)
            else:
                print(f"  ⚠️  Directory not found: {directory}")
        
        # Load structured files
        structured_files = data_sources.get('structured_files', [])
        for struct_file in structured_files:
            file_path = struct_file.get('file')
            if file_path and os.path.exists(file_path):
                print(f"  📊 Loading structured file: {file_path}")
                docs, metadata = self.load_structured_file(struct_file)
                self.documents.extend(docs)
                self.document_metadata.extend(metadata)
            else:
                print(f"  ⚠️  Structured file not found: {file_path}")
        
        # Load CSV files from custom config
        csv_files = data_sources.get('csv_files', [])
        for csv_file in csv_files:
            if os.path.exists(csv_file):
                print(f"  📊 Loading CSV: {csv_file}")
                docs, metadata = self.load_csv_content(csv_file)
                self.documents.extend(docs)
                self.document_metadata.extend(metadata)
            else:
                print(f"  ⚠️  CSV file not found: {csv_file}")
        
        # Process documents
        if self.documents:
            print(f"📋 Processing {len(self.documents)} raw documents...")
            processed_docs, processed_metadata = self.process_documents(
                self.documents, self.document_metadata
            )
            print(f"✅ Final document count: {len(processed_docs)}")
            
            # Convert to Document objects for RAGAS compatibility
            if LANGCHAIN_AVAILABLE:
                document_objects = []
                for doc_content, metadata in zip(processed_docs, processed_metadata):
                    doc_obj = Document(
                        page_content=doc_content,
                        metadata=metadata
                    )
                    document_objects.append(doc_obj)
                print(f"📄 Created {len(document_objects)} Document objects for RAGAS")
                return document_objects, processed_metadata
            else:
                print("⚠️ langchain not available, returning strings")
                return processed_docs, processed_metadata
        else:
            print("⚠️  No documents loaded. Using default document corpus.")
            default_docs, default_metadata = self.get_default_documents()
            
            # Convert default documents to Document objects
            if LANGCHAIN_AVAILABLE and default_docs and isinstance(default_docs[0], str):
                document_objects = []
                for doc_content, metadata in zip(default_docs, default_metadata):
                    doc_obj = Document(
                        page_content=doc_content,
                        metadata=metadata
                    )
                    document_objects.append(doc_obj)
                return document_objects, default_metadata
            else:
                return default_docs, default_metadata
    
    def load_pdf(self, file_path: str) -> Tuple[List[str], List[Dict]]:
        """Load and extract text from PDF file"""
        documents = []
        metadata = []
        
        if not PDF_AVAILABLE:
            print(f"  ❌ PDF processing not available for {file_path}")
            return documents, metadata
        
        try:
            # Try pdfplumber first (better text extraction)
            with pdfplumber.open(file_path) as pdf:
                full_text = ""
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        full_text += page_text + "\n\n"
                
                # Chunk the document
                chunks = self.chunk_text(full_text)
                for i, chunk in enumerate(chunks):
                    documents.append(chunk)
                    metadata.append({
                        'source_file': file_path,
                        'file_type': 'pdf',
                        'chunk_id': i,
                        'total_chunks': len(chunks)
                    })
        
        except Exception as e:
            print(f"  ❌ Error loading PDF {file_path}: {e}")
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    full_text = ""
                    
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            full_text += page_text + "\n\n"
                    
                    chunks = self.chunk_text(full_text)
                    for i, chunk in enumerate(chunks):
                        documents.append(chunk)
                        metadata.append({
                            'source_file': file_path,
                            'file_type': 'pdf',
                            'chunk_id': i,
                            'total_chunks': len(chunks)
                        })
                        
            except Exception as e2:
                print(f"  ❌ PyPDF2 fallback failed for {file_path}: {e2}")
        
        return documents, metadata
    
    def load_docx(self, file_path: str) -> Tuple[List[str], List[Dict]]:
        """Load and extract text from Word document"""
        documents = []
        metadata = []
        
        if not DOCX_AVAILABLE:
            print(f"  ❌ DOCX processing not available for {file_path}")
            return documents, metadata
        
        try:
            doc = DocxDocument(file_path)
            full_text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text += paragraph.text + "\n"
            
            chunks = self.chunk_text(full_text)
            for i, chunk in enumerate(chunks):
                documents.append(chunk)
                metadata.append({
                    'source_file': file_path,
                    'file_type': 'docx',
                    'chunk_id': i,
                    'total_chunks': len(chunks)
                })
                
        except Exception as e:
            print(f"  ❌ Error loading DOCX {file_path}: {e}")
        
        return documents, metadata
    
    def load_text_file(self, file_path: str) -> Tuple[List[str], List[Dict]]:
        """Load text from plain text file"""
        documents = []
        metadata = []
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                print(f"  ❌ Could not decode {file_path} with any encoding")
                return documents, metadata
            
            # Filter by language if enabled
            if self.processing_config.get('filter_language'):
                if not self.is_target_language(content):
                    print(f"  ⚠️  Skipping {file_path} - not target language")
                    return documents, metadata
            
            chunks = self.chunk_text(content)
            for i, chunk in enumerate(chunks):
                documents.append(chunk)
                metadata.append({
                    'source_file': file_path,
                    'file_type': 'txt',
                    'chunk_id': i,
                    'total_chunks': len(chunks)
                })
                
        except Exception as e:
            print(f"  ❌ Error loading text file {file_path}: {e}")
        
        return documents, metadata
    
    def load_structured_file(self, file_config: Dict) -> Tuple[List[str], List[Dict]]:
        """Load documents from structured files (CSV/Excel)"""
        documents = []
        metadata = []
        
        file_path = file_config.get('file')
        document_column = file_config.get('document_column', 'content')
        topic_column = file_config.get('topic_column', 'topic')
        
        try:
            # Load based on file extension
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
            elif file_path.endswith(('.xlsx', '.xls')):
                df = pd.read_excel(file_path)
            else:
                print(f"  ❌ Unsupported structured file format: {file_path}")
                return documents, metadata
            
            if document_column not in df.columns:
                print(f"  ❌ Column '{document_column}' not found in {file_path}")
                return documents, metadata
            
            for idx, row in df.iterrows():
                content = str(row[document_column])
                if pd.isna(content) or len(content.strip()) < self.processing_config.get('min_document_length', 50):
                    continue
                
                chunks = self.chunk_text(content)
                for i, chunk in enumerate(chunks):
                    documents.append(chunk)
                    metadata.append({
                        'source_file': file_path,
                        'file_type': 'structured',
                        'chunk_id': i,
                        'total_chunks': len(chunks),
                        'row_index': idx,
                        'topic': row.get(topic_column, 'unknown') if topic_column in df.columns else 'unknown'
                    })
                    
        except Exception as e:
            print(f"  ❌ Error loading structured file {file_path}: {e}")
        
        return documents, metadata
    
    def load_csv_content(self, file_path: str) -> Tuple[List[str], List[Dict]]:
        """Load content from CSV file with JSON content field"""
        documents = []
        metadata = []
        
        try:
            import json
            import re
            
            df = pd.read_csv(file_path)
            print(f"  📊 Loading CSV: {file_path} ({len(df)} rows)")
            
            # Check for required columns
            if 'id' not in df.columns or 'content' not in df.columns:
                print(f"  ❌ CSV file missing required columns 'id' or 'content': {file_path}")
                return documents, metadata
            
            # Get content_json_fields configuration
            csv_config = self.config.get('data_sources', {}).get('csv', {})
            format_config = csv_config.get('format', {})
            content_json_fields = format_config.get('content_json_fields', {})
            
            # Default field mapping if not configured
            text_field = content_json_fields.get('text', 'text')
            title_field = content_json_fields.get('title', 'title')
            metadata_field = content_json_fields.get('metadata', 'metadata')
            
            print(f"  🔧 JSON field mapping: text='{text_field}', title='{title_field}', metadata='{metadata_field}'")
            
            for idx, row in df.iterrows():
                try:
                    # Extract content from JSON field
                    content_json = row['content'] if 'content' in row and pd.notna(row['content']) else '{}'
                    if isinstance(content_json, str) and content_json.strip():
                        try:
                            content_data = json.loads(content_json)
                            
                            # Use configured field names with fallback to hardcoded names
                            if text_field and text_field != 'null':
                                text_content = content_data.get(text_field, '')
                                if not text_content:
                                    text_content = content_data.get('text', content_data.get('Content', ''))
                            else:
                                # If text_field is None/null, use direct content or fallback to Content
                                text_content = str(content_json) if not content_json.startswith('{') else content_data.get('Content', content_data.get('text', ''))
                            
                            if title_field and title_field != 'null':
                                content_title = content_data.get(title_field, '')
                                if not content_title:
                                    content_title = content_data.get('title', content_data.get('Title', ''))
                            else:
                                content_title = content_data.get('Title', content_data.get('title', ''))
                            
                            if metadata_field and metadata_field != 'null':
                                content_metadata = content_data.get(metadata_field, {})
                                if not isinstance(content_metadata, dict):
                                    content_metadata = {}
                            else:
                                content_metadata = {}
                            
                            # Extract language and other metadata
                            content_language = content_data.get('language', 'EN')
                            if isinstance(content_metadata, dict) and content_language == 'EN':
                                content_language = content_metadata.get('language', 'EN')
                            
                            content_source = content_data.get('source', '')
                            if isinstance(content_metadata, dict) and not content_source:
                                content_source = content_metadata.get('source', '')
                            
                        except json.JSONDecodeError:
                            # Fallback to raw string if not JSON
                            text_content = str(content_json)
                            content_metadata = {}
                            content_language = 'EN'
                            content_title = ''
                            content_source = ''
                    else:
                        text_content = str(content_json) if content_json else ''
                        content_metadata = {}
                        content_language = 'EN'
                        content_title = ''
                        content_source = ''
                    
                    # Skip if content is too short
                    if len(text_content.strip()) < 10:  # Minimum content check
                        continue
                    
                    # Clean the text
                    cleaned_text = self.clean_text(text_content)
                    
                    if len(cleaned_text) >= 10:
                        # Enhanced content for RAGAS compatibility
                        enhanced_content = self._enhance_content_for_ragas(cleaned_text, content_data)
                        
                        # For CSV: Each row = 1 document chunk (1:1 mapping as requested)
                        documents.append(enhanced_content)
                        
                        # Enhanced metadata with RAGAS-required attributes
                        enhanced_metadata = {
                            'source_file': file_path,
                            'file_type': 'csv',
                            'csv_id': row['id'] if 'id' in row and pd.notna(row['id']) else idx,
                            'csv_row': idx,
                            'template_key': row['template_key'] if 'template_key' in row and pd.notna(row['template_key']) else '',
                            'source': row['source'] if 'source' in row and pd.notna(row['source']) else '',
                            'author': row['author'] if 'author' in row and pd.notna(row['author']) else '',
                            'created_at': row['created_at'] if 'created_at' in row and pd.notna(row['created_at']) else '',
                            'updated_at': row['updated_at'] if 'updated_at' in row and pd.notna(row['updated_at']) else '',
                            'content_title': content_title,
                            'content_source': content_source,
                            'content_language': content_language,
                            'content_metadata': content_metadata,
                            'word_count': len(cleaned_text.split()),
                            'chunk_id': 0,  # Each CSV row is one chunk
                            'total_chunks': 1,
                            # RAGAS-required attributes for knowledge graph compatibility
                            'document_id': f"doc_{idx}",
                            'entities': self._extract_entities(cleaned_text),
                            'keyphrases': self._extract_keyphrases(cleaned_text),
                            'summary': self._create_summary(cleaned_text),
                            'themes': self._extract_themes(cleaned_text),
                            'headlines': self._extract_headlines(cleaned_text),
                            'has_technical_content': self._has_technical_content(cleaned_text),
                            'page_content': enhanced_content  # RAGAS expects page_content
                        }
                        metadata.append(enhanced_metadata)
                
                except Exception as e:
                    print(f"  ⚠️  Error processing CSV row {idx}: {e}")
                    continue
            
            print(f"  ✅ Loaded {len(documents)} content chunks from CSV")
            
        except Exception as e:
            print(f"  ❌ Error loading CSV {file_path}: {e}")
        
        return documents, metadata
    
    def _enhance_content_for_ragas(self, text_content: str, content_data: Dict) -> str:
        """Enhance content for RAGAS compatibility"""
        # Ensure minimum content length
        if len(text_content.strip()) < 50:
            title = content_data.get('title', 'Document')
            source = content_data.get('source', 'Unknown')
            enhanced_content = f"Title: {title}\n\nContent: {text_content}\n\nSource: {source}"
            return enhanced_content
        return text_content
    
    def _extract_entities(self, content: str) -> list:
        """Extract entities from content"""
        import re
        entities = []
        
        # Capitalized words (proper nouns)
        capitalized = re.findall(r'\b[A-Z][a-z]+\b', content)
        entities.extend(capitalized[:8])
        
        # Technical terms/abbreviations  
        technical = re.findall(r'\b[A-Z]{2,}\b', content)
        entities.extend(technical[:5])
        
        # Numbers and codes
        numbers = re.findall(r'\b\d+(?:\.\d+)?\b', content)
        entities.extend(numbers[:3])
        
        # Remove duplicates while preserving order
        return list(dict.fromkeys(entities))
    
    def _extract_keyphrases(self, content: str) -> list:
        """Extract keyphrases from content"""
        import re
        words = content.split()
        keyphrases = []
        
        for word in words:
            # Clean word
            clean_word = re.sub(r'[^\w]', '', word.lower())
            
            # Include words longer than 3 characters
            if len(clean_word) > 3 and clean_word.isalpha():
                keyphrases.append(clean_word)
        
        # Return unique keyphrases
        return list(dict.fromkeys(keyphrases))[:15]
    
    def _create_summary(self, content: str) -> str:
        """Create a summary of content"""
        sentences = content.split('.')[:3]  # First 3 sentences
        summary = '. '.join(s.strip() for s in sentences if s.strip())
        return summary if summary else content[:200]
    
    def _extract_themes(self, content: str) -> list:
        """Extract themes from content"""
        themes = []
        technical_terms = ['system', 'process', 'error', 'code', 'function', 
                          'operation', 'performance', 'quality', 'standard', 'procedure']
        
        content_lower = content.lower()
        for term in technical_terms:
            if term in content_lower:
                themes.append(term)
        
        return themes[:5]
    
    def _extract_headlines(self, content: str) -> list:
        """Extract headlines from content"""
        sentences = [s.strip() for s in content.split('.') if len(s.strip()) > 10]
        headlines = []
        
        for sentence in sentences[:5]:
            if len(sentence) <= 80:
                headlines.append(sentence)
            else:
                headlines.append(sentence[:77] + "...")
        
        return headlines
    
    def _has_technical_content(self, content: str) -> bool:
        """Check if content has technical information"""
        technical_indicators = ['error', 'code', 'system', 'function', 'process',
                              'procedure', 'standard', 'specification', 'parameter']
        content_lower = content.lower()
        return any(indicator in content_lower for indicator in technical_indicators)
    
    def load_directory(self, directory_path: str) -> Tuple[List[str], List[Dict]]:
        """Load all supported files from directory recursively"""
        all_documents = []
        all_metadata = []
        
        supported_extensions = {
            '.pdf': self.load_pdf,
            '.txt': self.load_text_file,
            '.docx': self.load_docx,
            '.doc': self.load_docx,
        }
        
        print(f"    🔍 Scanning directory: {directory_path}")
        
        for root, dirs, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                file_ext = Path(file).suffix.lower()
                
                if file_ext in supported_extensions:
                    print(f"      📄 Processing: {file}")
                    try:
                        load_function = supported_extensions[file_ext]
                        docs, metadata = load_function(file_path)
                        all_documents.extend(docs)
                        all_metadata.extend(metadata)
                    except Exception as e:
                        print(f"      ❌ Error processing {file}: {e}")
        
        print(f"    ✅ Loaded {len(all_documents)} documents from directory")
        return all_documents, all_metadata
    
    def chunk_text(self, text: str) -> List[str]:
        """Split text into manageable chunks"""
        chunk_size = self.processing_config.get('chunk_size', 1000)
        chunk_overlap = self.processing_config.get('chunk_overlap', 200)
        min_chunk_size = self.processing_config.get('min_chunk_size', 100)
        
        if len(text) <= chunk_size:
            return [text.strip()] if text.strip() else []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Find a good break point (sentence end)
            if end < len(text):
                # Look for sentence breaks near the end
                for i in range(min(50, chunk_size // 4)):
                    if text[end - i] in '.!?':
                        end = end - i + 1
                        break
            
            chunk = text[start:end].strip()
            
            if len(chunk) >= min_chunk_size:
                chunks.append(chunk)
            
            start = end - chunk_overlap
            
            if start >= len(text):
                break
        
        return chunks
    
    def process_documents(self, documents: List[str], metadata: List[Dict]) -> Tuple[List[str], List[Dict]]:
        """Post-process loaded documents"""
        processed_docs = []
        processed_metadata = []
        
        # Use different minimum lengths for different file types
        min_length = self.processing_config.get('min_document_length', 100)
        
        for doc, meta in zip(documents, metadata):
            # Clean and validate document
            cleaned_doc = self.clean_text(doc)
            
            # For CSV files, use a lower minimum length since each row is intentionally a chunk
            effective_min_length = min_length
            if meta.get('file_type') == 'csv':
                effective_min_length = 10  # Lower threshold for CSV content
            
            if len(cleaned_doc) >= effective_min_length:
                processed_docs.append(cleaned_doc)
                processed_metadata.append(meta)
            else:
                print(f"  ⚠️  Skipping short document: {len(cleaned_doc)} chars (min: {effective_min_length})")
        
        print(f"  📋 Filtered documents: {len(documents)} → {len(processed_docs)}")
        return processed_docs, processed_metadata
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        import re
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters if configured
        if self.processing_config.get('remove_special_chars', False):
            text = re.sub(r'[^\w\s\.\,\!\?\;\:]', '', text)
        
        # Strip and return
        return text.strip()
    
    def is_target_language(self, text: str) -> bool:
        """Check if text is in target language"""
        if not LANGDETECT_AVAILABLE:
            return True
        
        target_lang = self.processing_config.get('filter_language')
        if not target_lang:
            return True
        
        try:
            detected_lang = detect(text)
            return detected_lang == target_lang
        except:
            return True  # Default to including if detection fails
    
    def extract_topics_from_text(self, text: str, n_topics: int = 5) -> List[str]:
        """Extract topics from text using available methods"""
        topics = []
        
        # Try KeyBERT first
        if KEYBERT_AVAILABLE and hasattr(self, 'keybert_model'):
            try:
                keywords = self.keybert_model.extract_keywords(
                    text, 
                    keyphrase_ngram_range=(1, 2), 
                    stop_words='english'
                )
                topics.extend([kw[0] for kw in keywords[:n_topics]])
            except Exception as e:
                print(f"  ⚠️  KeyBERT topic extraction failed: {e}")
        
        # Try YAKE if KeyBERT failed
        if not topics and YAKE_AVAILABLE and hasattr(self, 'yake_extractor'):
            try:
                keywords = self.yake_extractor.extract_keywords(text)
                topics.extend([kw[1] for kw in keywords[:n_topics]])
            except Exception as e:
                print(f"  ⚠️  YAKE topic extraction failed: {e}")
        
        # Fallback to simple extraction
        if not topics:
            import re
            words = re.findall(r'\b[a-zA-Z]{4,}\b', text.lower())
            word_freq = {}
            for word in words:
                word_freq[word] = word_freq.get(word, 0) + 1
            
            # Get most frequent words as topics
            sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
            topics = [word for word, freq in sorted_words[:n_topics] if freq > 1]
        
        return topics[:n_topics]

    def get_document_stats(self) -> Dict:
        """Get statistics about loaded documents"""
        if not self.documents:
            return {}
        
        # Group by source file
        source_files = {}
        for meta in self.document_metadata:
            source = meta.get('source_file', 'unknown')
            if source not in source_files:
                source_files[source] = {'chunks': 0, 'file_type': meta.get('file_type', 'unknown')}
            source_files[source]['chunks'] += 1
        
        stats = {
            'total_documents': len(self.documents),
            'total_source_files': len(source_files),
            'files_by_type': {},
            'chunks_per_file': {},
            'avg_document_length': sum(len(doc) for doc in self.documents) / len(self.documents)
        }
        
        # Calculate file type statistics
        for source, info in source_files.items():
            file_type = info['file_type']
            if file_type not in stats['files_by_type']:
                stats['files_by_type'][file_type] = {'count': 0, 'total_chunks': 0}
            
            stats['files_by_type'][file_type]['count'] += 1
            stats['files_by_type'][file_type]['total_chunks'] += info['chunks']
            stats['chunks_per_file'][source] = info['chunks']
        
        return stats
    
    def get_default_documents(self) -> Tuple[List[str], List[Dict]]:
        """Return default documents when no custom data is loaded"""
        default_docs = [
            "Type 2 diabetes is a chronic condition affecting glucose metabolism. Insulin resistance and beta-cell dysfunction are key pathophysiological mechanisms.",
            "Hypertension diagnosis requires multiple blood pressure measurements. Systolic >140 mmHg or diastolic >90 mmHg indicates hypertension.",
            "Pneumonia involves lung inflammation caused by bacterial, viral, or fungal pathogens. Streptococcus pneumoniae is the most common bacterial cause.",
            "Cardiovascular disease risk factors include modifiable elements like smoking, diet, exercise, and non-modifiable factors like age and genetics.",
            "Asthma management involves bronchodilators for acute symptoms and anti-inflammatory medications for long-term control.",
            "Machine learning algorithms learn patterns from data without explicit programming. Supervised, unsupervised, and reinforcement learning are main paradigms.",
            "Cloud computing provides scalable computing resources through virtualization and distributed systems across multiple data centers.",
            "Cybersecurity protects digital assets through threat detection, access controls, encryption, and incident response protocols.",
            "Blockchain technology uses cryptographic hashing and distributed consensus mechanisms to create immutable transaction records.",
            "API integration enables software interoperability through standardized communication protocols and data exchange formats."
        ]
        
        default_metadata = [
            {'source_file': 'default', 'file_type': 'builtin', 'chunk_id': i, 'total_chunks': len(default_docs)}
            for i in range(len(default_docs))
        ]
        
        return default_docs, default_metadata
    
    def get_topics_from_metadata(self) -> List[str]:
        """Extract unique topics from loaded document metadata"""
        topics = set()
        
        for meta in self.document_metadata:
            # From structured file topics
            if 'topic' in meta:
                topics.add(meta['topic'])
            
            # From extracted topics
            if 'extracted_topics' in meta:
                topics.update(meta['extracted_topics'])
        
        # Convert to list and add some generic topics if none found
        topic_list = list(topics)
        if not topic_list:
            topic_list = [
                "concept", "method", "approach", "technique", "strategy",
                "system", "process", "framework", "model", "solution"
            ]
        
        return topic_list
    
    def print_loading_summary(self):
        """Print summary of loaded documents"""
        if not self.documents:
            print("📋 No custom documents loaded - using default corpus")
            return
        
        print(f"\n📋 Document Loading Summary:")
        print("=" * 40)
        print(f"Total documents loaded: {len(self.documents)}")
        
        # Group by file type
        file_types = {}
        for meta in self.document_metadata:
            file_type = meta.get('file_type', 'unknown')
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        for file_type, count in file_types.items():
            print(f"  {file_type.upper()}: {count} chunks")
        
        # Show topics if available
        topics = self.get_topics_from_metadata()
        if topics:
            print(f"\nExtracted topics: {', '.join(topics[:5])}")
            if len(topics) > 5:
                print(f"  ... and {len(topics) - 5} more")

    def _extract_themes(self, text):
        """Extract main themes from text"""
        try:
            words = text.lower().split()
            # Simple theme extraction based on common patterns
            themes = []
            
            # Technical themes
            tech_keywords = ['api', 'database', 'system', 'framework', 'technology', 'software', 'hardware']
            if any(keyword in words for keyword in tech_keywords):
                themes.append('technical')
            
            # Business themes
            business_keywords = ['business', 'market', 'customer', 'revenue', 'strategy', 'management']
            if any(keyword in words for keyword in business_keywords):
                themes.append('business')
            
            # Educational themes
            edu_keywords = ['learn', 'education', 'training', 'course', 'study', 'knowledge']
            if any(keyword in words for keyword in edu_keywords):
                themes.append('educational')
            
            return themes if themes else ['general']
        except Exception:
            return ['general']

    def _extract_headlines(self, text):
        """Extract potential headlines from text"""
        try:
            lines = text.split('\n')
            headlines = []
            
            for line in lines[:5]:  # Check first 5 lines
                line = line.strip()
                if line and (len(line) < 100) and (line.endswith(':') or line.isupper() or any(char in line for char in ['#', '*', '-'])):
                    headlines.append(line)
            
            return headlines if headlines else [text.split('.')[0][:50] + '...']
        except Exception:
            return [text[:50] + '...']

    def _has_technical_content(self, text):
        """Check if text contains technical content"""
        try:
            tech_indicators = [
                'api', 'database', 'sql', 'json', 'xml', 'http', 'https',
                'function', 'method', 'class', 'variable', 'algorithm',
                'framework', 'library', 'module', 'interface', 'protocol'
            ]
            text_lower = text.lower()
            return any(indicator in text_lower for indicator in tech_indicators)
        except Exception:
            return False

def main():
    """Test document loading functionality"""
    # Example configuration
    config = {
        'custom_data': {
            'enabled': True,
            'data_sources': {
                'pdf_files': ['example.pdf'],
                'text_files': ['example.txt'],
                'directories': ['documents/']
            },
            'processing': {
                'chunk_size': 1000,
                'chunk_overlap': 200,
                'filter_by_language': True,
                'target_language': 'en'
            },
            'topic_extraction': {
                'enabled': True,
                'method': 'keybert',
                'max_topics_per_document': 3
            }
        }
    }
    
    loader = DocumentLoader(config)
    documents, metadata = loader.load_all_documents()
    loader.print_loading_summary()
    
    print(f"\nSample document:\n{documents[0][:200]}..." if documents else "No documents loaded")

if __name__ == "__main__":
    main()
